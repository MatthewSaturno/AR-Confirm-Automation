#Disclaimer: Some of the below code has been copied from StackOverflow and other online posts and customized to render\
#the program.

import csv
import PyPDF2
import docx
from docx2pdf import convert
from tabulate import tabulate
import pandas
import pandas as pd
import numpy as np

#Extract from the customer csv
customer_file = input("Please paste customer file name here \n")
customer_data = open(customer_file,encoding='utf-8-sig')
csv_data = csv.reader(customer_data)
customer_lines = list(csv_data)

#Extract from the invoice csv
invoice_file = input("\nPlease paste invoice file name here \n")
invoice_data = open(invoice_file,encoding='utf-8-sig')
csv_data = csv.reader(invoice_data)
invoice_lines = list(csv_data)

#Input info into text doc
def create_text_doc(row, ey_person, ey_email, date):
    
    customer = row[0]
    customer_address = row[1]
    city = row[2]
    province = row[3]
    postal = row[4]
    main_person = row[5]
    contact_email = row[6]
    company = row[7]
    company_contact = row[9]
    contact_title = row[10] 
   

    text1=date+'''

Attn: ''' + main_person+'\n'\
+customer+'\n'\
+customer_address+'\n'\
+city+', '+province+' '+postal+'\n'*2\
+'''Dear ''' + main_person +','+ '''

Our auditors, COMPANY NAME HERE (Attention: '''+ey_person+''', ADDRESS HERE)'''\
+''', are auditing our financial statements and wish to obtain direct confirmation of the amount owed to us as of the date'''\
+''' indicated below. Please compare the balance shown below with your records as of the date indicated and note the details'''\
+''' of exceptions (if any) in the space provided below or on an attachment. Then please sign this request in the space provided'''\
+''' and email a scanned copy of your response to '''+ey_email+'''

This is not a request for payment and remittances should not be sent to COMPANY NAME HERE.

Confirmation date: '''+confirm_date
    text2='''
Very truly yours,



'''\
+company_contact+', '+ contact_title+'\n'\
+company+'''


COMPANY NAME HERE:  

The account balance shown above is correct as of the date indicated, except for:

______________________________________________________________________
______________________________________________________________________

Signed ______________________________
Title  ________________________________
Date   ______________________________
'''

    return text1, text2, customer

#Grab invoice data
def grab_invoice_data(customer):
    a_list = []
    for item in invoice_lines[1:]:
        if item[0] == customer:
            a_list.append(item[1:])
            
    return a_list

#Save to Word and style
def savetoword(doc_count,text):
    doc1 = docx.Document()
    font = doc1.styles['Normal'].font
    font.name = 'Arial'
    
    #Add first text
    doc1.add_paragraph(text[0])
    
    #Add table
    add_a_table(data,count,headers,doc1)
    
    #Add second text
    doc1.add_paragraph(text[1])
    
    try:
        doc1.save(f"C:\\FILE-PATH-DESTINATION-HERE\\{doc_count}-{text[2]}.docx")
    except:
        doc1.save(f"C:\\FILE-PATH-DESTINATION-HERE\\{doc_count}-naming-error.docx")

#Add the table of invoices    
def add_a_table(data,count,headers,doc):

    df = pandas.DataFrame(data,count,headers)
    # add a table to the end and create a reference variable
    # extra row is so we can add the header row
    t = doc.add_table(df.shape[0]+1, df.shape[1])
    t.style='Table Grid'

    # add the header rows.
    for j in range(df.shape[-1]):
        t.cell(0,j).text = str(df.columns[j])
    for cell in t.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])

            
#Make a new function to convert to PDF
#TAKEN OUT TEMPORARILY BC IT'S GLITCHY
#Save to PDF
# convert("document.docx")
# convert("my_docx_folder/")

#RUN
doc_count = 1
#Enter in common information 
ey_person = input("\nPlease enter the name of the employee contact \n")
ey_email = input("\nPlease enter the email of the employee of contact \n")
date = input("\nWhat is the date of the letter? \n")
confirm_date = input("\nWhat is the confirmation date? \n")


#Create text of confirmations
for item in customer_lines[1:]:
    a = create_text_doc(item,ey_person, ey_email, date)
    #Make table of invoice data
    data = grab_invoice_data(item[0])
    count = list(range(len(data)))
    headers = ["Invoice #", "Invoice Date", "Invoice Amount"]
    
    #Save to word
    savetoword(doc_count,a)
    doc_count+=1


    
    
